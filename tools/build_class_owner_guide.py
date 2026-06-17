from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Cpp_Class_Structure_Guide.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.8)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 21, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(9.8)
        style.paragraph_format.space_after = Pt(2)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 70, "bottom": 70, "start": 100, "end": 100}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("ShowDown 클래스 담당자 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("C++ 중심 구조에서 누가 어떤 클래스를 만지는지 쉽게 보는 문서").bold = True

    doc.add_heading("1. 핵심 원칙", level=1)
    bullets(doc, [
        "코어 클래스는 박형빈님이 주로 만집니다.",
        "연출/사물형 UI 클래스는 임현진과 이진헌이 같이 만집니다.",
        "LLM, DB, 멀티 클래스는 이진헌이 주로 만집니다.",
        "맵/레벨 통합은 임현진이 관리합니다.",
        "이벤트 이름이나 공용 함수 이름을 바꿀 때는 반드시 팀에 공유합니다.",
    ])

    doc.add_heading("2. 전체 작업 흐름", level=1)
    bullets(doc, [
        "사물 클릭 -> PlayerController에 요청.",
        "PlayerController -> GameMode/코어 시스템에 요청.",
        "코어 시스템 -> 규칙 판단.",
        "GameState -> 이벤트 발생.",
        "PresentationDirector -> 이벤트를 받아 연출 실행.",
    ])

    doc.add_heading("3. 1차 필수 클래스", level=1)
    doc.add_paragraph("먼저 만들어야 하는 클래스입니다. 이게 되어야 게임 한 판이 돌아갑니다.")
    add_table(doc, ["클래스", "주 담당", "같이 확인", "역할 / 주의점"], [
        ["ASDGameMode", "박형빈", "이진헌", "게임 시작, 라운드 진행, 승패/룰렛 최종 판단."],
        ["ASDGameState", "박형빈", "임현진, 이진헌", "현재 Phase, 베팅, 목숨 상태 보관. 연출팀이 받을 이벤트를 여기서 발생."],
        ["ASDPlayerController", "박형빈 + 이진헌", "임현진", "입력/클릭을 받아 코어에 요청. 직접 승패 계산 금지."],
        ["ASDPlayerState", "박형빈", "이진헌", "플레이어별 목숨, 베팅, 카드 정보. 멀티 때 숨길 정보 주의."],
        ["USDCardSystemComponent", "박형빈", "이진헌", "덱 생성, 셔플, 손패 지급, 카드 폐기."],
        ["USDBettingSystemComponent", "박형빈", "임현진, 이진헌", "Check/Raise/Call/Fold 처리. 베팅 변경 이벤트 필요."],
        ["USDRoundResolverComponent", "박형빈", "임현진, 이진헌", "카드 공개, 승패 판정, 동점, 7 폴드 예외."],
        ["USDRouletteSystemComponent", "박형빈", "임현진, 이진헌", "장전 수 / 6 확률로 명중/불발 계산. 연출은 직접 실행하지 않음."],
        ["USDBossAIComponent", "박형빈", "이진헌", "보스의 기본 베팅/콜/폴드 판단. LLM이 판단하지 않음."],
    ])

    doc.add_heading("4. 2차 연출 / 사물형 UI 클래스", level=1)
    doc.add_paragraph("코어 루프와 동시에 테스트맵에서 만들어도 됩니다. 임현진과 이진헌이 계속 상의하면서 같이 만집니다.")
    add_table(doc, ["클래스", "주 담당", "같이 확인", "역할 / 주의점"], [
        ["ASDInteractableActor", "임현진 + 이진헌", "박형빈", "모든 클릭 가능한 사물의 부모. Hover, Unhover, Interact."],
        ["ASDDiegeticButtonActor", "임현진 + 이진헌", "박형빈", "Call/Raise/Fold/Check 물리 버튼. 클릭 시 코어 함수 요청."],
        ["ASDBulletSelectorActor", "임현진 + 이진헌", "박형빈", "1~6발 선택 장치. RaiseTo(BulletCount) 요청."],
        ["ASDCardActor", "임현진 + 이진헌", "박형빈", "월드 카드 오브젝트. 카드 선택/부착/공개 표시."],
        ["ASDLifeCardActor", "임현진 + 이진헌", "박형빈", "목숨 1개를 나타내는 생명 카드."],
        ["ASDCRTMonitorActor", "임현진 + 이진헌", "박형빈", "게임 안 모니터에 상태/경고 표시."],
        ["ASDPresentationDirector", "임현진 + 이진헌", "박형빈", "GameState 이벤트를 받아 연출 실행. 입력 잠금/해제 관리."],
        ["ASDRouletteSequenceActor", "임현진 + 이진헌", "박형빈", "장전, 약실 회전, 방아쇠, 명중/불발 연출."],
        ["ASDCardRevealSequenceActor", "임현진 + 이진헌", "박형빈", "카드 뒤집기, 줌인, 공개 사운드."],
        ["ASDBettingPresentationActor", "임현진 + 이진헌", "박형빈", "총알 놓기, Raise 위험 강조."],
        ["ASDStageMoodController", "임현진 + 이진헌", "김윤아", "Stage별 조명, 노이즈, 분위기 변화."],
    ])

    doc.add_heading("5. 3차 기술 어필 클래스", level=1)
    doc.add_paragraph("평가 때 기술적으로 보여주기 위한 기능입니다. 처음부터 본 게임에 붙이지 말고 테스트맵에서 먼저 성공시킵니다.")
    add_table(doc, ["클래스", "주 담당", "같이 확인", "역할 / 주의점"], [
        ["USDLLMSubsystem", "이진헌", "임현진", "보스 대사 API 요청/응답/실패 처리. API Key 커밋 금지."],
        ["USDBossDialogueComponent", "이진헌", "김윤아, 임현진", "대사 풀 선택, LLM 요청, 실패 시 기본 대사 출력."],
        ["USDShowDownSaveGame", "이진헌", "박형빈", "기억 조각, 스킨, 승패 기록 로컬 저장."],
        ["USDSaveSubsystem", "이진헌", "박형빈", "저장/불러오기 관리."],
        ["USDDatabaseSubsystem", "이진헌", "박형빈", "외부 DB 기록 저장/불러오기. 비밀번호 커밋 금지."],
        ["ASDMultiplayerGameMode", "이진헌", "박형빈", "멀티 1:1 게임 시작과 서버 권위 처리."],
        ["ASDMultiplayerGameState", "이진헌", "박형빈", "멀티 Phase, 베팅, 목숨, 결과 복제."],
        ["ASDMultiplayerPlayerController", "이진헌", "박형빈", "Server_SelectCard, Server_RaiseTo 같은 멀티 요청."],
    ])

    doc.add_heading("6. 김윤아 담당", level=1)
    bullets(doc, [
        "김윤아는 C++ 클래스보다는 에셋, 맵, 소품, 머티리얼, 스토리 오브젝트를 담당합니다.",
        "에셋 출처와 라이선스를 정리합니다.",
        "Stage별 방 변화에 필요한 소품을 준비합니다.",
        "메인 레벨에 넣을 에셋은 임현진과 상의해서 반영합니다.",
    ])

    doc.add_heading("7. 공통 이벤트", level=1)
    add_table(doc, ["이벤트", "발생 담당", "사용 담당"], [
        ["OnPhaseChanged(NewPhase)", "박형빈", "임현진 + 이진헌"],
        ["OnBetChanged(Side, BulletCount)", "박형빈", "임현진 + 이진헌"],
        ["OnCardsRevealed(PlayerCard, BossCard)", "박형빈", "임현진 + 이진헌"],
        ["OnRoundResolved(Result)", "박형빈", "임현진 + 이진헌"],
        ["OnRouletteStarted(Target, BulletCount)", "박형빈", "임현진 + 이진헌"],
        ["OnRouletteResult(Target, bHit)", "박형빈", "임현진 + 이진헌"],
        ["OnLifeChanged(Target, Life)", "박형빈", "임현진 + 이진헌"],
        ["OnStageChanged(Stage)", "박형빈", "임현진 + 이진헌 + 김윤아"],
        ["OnGameOver(Winner)", "박형빈", "임현진 + 이진헌"],
    ])

    doc.add_heading("8. 꼭 지킬 것", level=1)
    bullets(doc, [
        "코어 코드에서 카메라, 사운드, VFX를 직접 실행하지 않습니다.",
        "연출 코드에서 승패, 확률, 목숨 계산을 하지 않습니다.",
        "LLM은 보스 판단을 하지 않습니다. 대사만 만듭니다.",
        "멀티에서는 숨겨야 하는 카드 정보를 클라이언트에 함부로 복제하지 않습니다.",
        "클래스 이름, 함수 이름, 이벤트 이름을 바꾸기 전에는 팀에 공유합니다.",
        "테스트맵에서 성공한 뒤 본 게임에 연결합니다.",
    ])

    doc.add_heading("9. 한 줄 정리", level=1)
    doc.add_paragraph(
        "박형빈님은 코어 판단, 임현진과 이진헌은 연출/사물형 UI, 이진헌은 LLM/DB/멀티, 김윤아는 에셋/맵을 맡고, 서로는 GameState 이벤트로 연결합니다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Class Owner Guide")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

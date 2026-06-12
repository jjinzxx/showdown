from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Cpp_Class_Structure_Guide.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("ShowDown C++ 클래스 구조 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("Blueprint는 최소화하고, C++ 중심으로 코어/상호작용/연출/기술 기능을 나누는 기준 문서").bold = True

    doc.add_heading("1. 핵심 원칙", level=1)
    bullets(doc, [
        "코어는 게임 규칙만 처리한다.",
        "사물형 UI는 입력을 받아 코어 함수에 요청만 보낸다.",
        "연출은 코어 이벤트를 받아 화면, 사운드, 카메라로 보여준다.",
        "LLM, DB, 멀티는 테스트 레벨에서 먼저 성공시킨 뒤 본 게임에 연결한다.",
        "동작은 C++로 만들고, 에셋/사운드/수치만 에디터에서 조정한다.",
    ])

    doc.add_heading("2. 추천 폴더 구조", level=1)
    bullets(doc, [
        "Source/ShowDown/Public/Core, Private/Core: 게임 규칙.",
        "Source/ShowDown/Public/Data, Private/Data: enum, struct, DataAsset.",
        "Source/ShowDown/Public/Interaction, Private/Interaction: 클릭 가능한 사물.",
        "Source/ShowDown/Public/Presentation, Private/Presentation: 연출 제어.",
        "Source/ShowDown/Public/Boss, Private/Boss: 보스 AI와 대사.",
        "Source/ShowDown/Public/LLM, Save, Network: LLM, 저장/DB, 멀티.",
        "Content/ShowDown/Maps/Test: 테스트맵.",
        "Content/ShowDown/Art, VFX, Audio, Data: 에셋과 데이터.",
    ])

    doc.add_heading("3. 기본 데이터", level=1)
    add_table(doc, ["이름", "용도"], [
        ["ESDGamePhase", "DealCards, SelectCard, Betting, Reveal, Roulette, RoundEnd, GameOver 같은 진행 단계."],
        ["ESDPlayerSide", "Player, Boss, RemotePlayer 구분."],
        ["ESDActionType", "Check, Raise, Call, Fold, SelectCard 구분."],
        ["ESDRoundResult", "PlayerWin, BossWin, Draw, PlayerFold, BossFold."],
        ["FSDCard", "CardId, Number, bUsed를 가진 카드 데이터."],
        ["FSDBetState", "PlayerBet, BossBet, CurrentMaxBet 같은 베팅 상태."],
        ["FSDRoundState", "한 판의 카드, 베팅, 결과, 룰렛 정보를 묶는 상태."],
    ])

    doc.add_heading("4. 코어 게임 클래스", level=1)
    add_table(doc, ["클래스", "역할", "담당"], [
        ["ASDGameMode", "게임 시작, 라운드 진행, 승패/룰렛 최종 결정.", "박형빈"],
        ["ASDGameState", "현재 Phase, Stage, 베팅, 목숨 상태와 이벤트 보관.", "박형빈"],
        ["ASDPlayerController", "클릭/입력을 받아 코어에 요청.", "박형빈 + 이진헌"],
        ["ASDPlayerState", "플레이어별 목숨, 베팅, 손패 정보.", "박형빈"],
        ["USDCardSystemComponent", "덱 생성, 셔플, 손패 지급, 카드 폐기.", "박형빈"],
        ["USDBettingSystemComponent", "Check, Raise, Call, Fold 처리.", "박형빈"],
        ["USDRoundResolverComponent", "카드 공개, 승패 판정, 7 폴드 예외.", "박형빈"],
        ["USDRouletteSystemComponent", "장전 수 / 6 확률로 Hit/Miss 결정.", "박형빈"],
        ["USDBossAIComponent", "보스의 기본 베팅/콜/폴드 판단.", "박형빈"],
    ])

    doc.add_heading("5. 사물형 UI / 상호작용 클래스", level=1)
    add_table(doc, ["클래스", "역할", "담당"], [
        ["ASDInteractableActor", "모든 클릭 가능한 사물의 부모. Hover, Unhover, Interact.", "임현진 + 이진헌"],
        ["ASDDiegeticButtonActor", "Call/Raise/Fold/Check 물리 버튼.", "임현진 + 이진헌"],
        ["ASDBulletSelectorActor", "1~6발 선택 장치. RaiseTo 요청.", "임현진 + 이진헌"],
        ["ASDCardActor", "월드 카드 오브젝트. 선택, 부착, 공개 표시.", "임현진 + 이진헌"],
        ["ASDLifeCardActor", "목숨 1개를 나타내는 생명 카드.", "임현진 + 이진헌"],
        ["ASDCRTMonitorActor", "게임 안 모니터에 상태/경고 표시.", "임현진 + 이진헌"],
    ])

    doc.add_heading("6. 연출 클래스", level=1)
    add_table(doc, ["클래스", "역할", "담당"], [
        ["ASDPresentationDirector", "GameState 이벤트를 받아 연출 실행. 입력 잠금/해제 관리.", "임현진 + 이진헌"],
        ["ASDRouletteSequenceActor", "장전, 약실 회전, 방아쇠, 명중/불발 연출.", "임현진 + 이진헌"],
        ["ASDCardRevealSequenceActor", "카드 뒤집기, 줌인, 공개 사운드.", "임현진 + 이진헌"],
        ["ASDBettingPresentationActor", "총알 놓기, Raise 위험 강조.", "임현진 + 이진헌"],
        ["ASDStageMoodController", "Stage별 조명, 노이즈, Post Process 변화.", "임현진 + 이진헌"],
    ])

    doc.add_heading("7. LLM / 저장 / 멀티 클래스", level=1)
    add_table(doc, ["클래스", "역할", "담당"], [
        ["USDLLMSubsystem", "보스 대사 API 요청, 응답, 실패 처리.", "이진헌"],
        ["USDBossDialogueComponent", "대사 풀 선택, LLM 요청, 폴백 대사 출력.", "이진헌"],
        ["USDShowDownSaveGame", "기억 조각, 스킨, 승패 기록 로컬 저장.", "이진헌"],
        ["USDSaveSubsystem", "저장/불러오기 관리.", "이진헌"],
        ["USDDatabaseSubsystem", "외부 DB 기록 저장/불러오기.", "이진헌"],
        ["ASDMultiplayerGameMode", "멀티 1:1 게임 시작과 서버 권위 처리.", "이진헌"],
        ["ASDMultiplayerGameState", "멀티 Phase, 베팅, 목숨, 결과 복제.", "이진헌"],
        ["ASDMultiplayerPlayerController", "Server_SelectCard, Server_RaiseTo 같은 멀티 요청.", "이진헌"],
    ])

    doc.add_heading("8. DataAsset", level=1)
    add_table(doc, ["DataAsset", "담는 내용"], [
        ["USDStageConfigData", "Stage 번호, 목숨, 최소 베팅, AI 공격성, 노이즈 강도."],
        ["USDBossAIProfileData", "블러핑 확률, 폴드 확률, 위험 감수 성향."],
        ["USDDialoguePoolData", "Stage별 대사, Raise/Fold/Roulette 상황별 대사."],
        ["USDSkinData", "스킨 ID, 이름, 가격, 머티리얼/메시, 해금 여부."],
    ])

    doc.add_heading("9. 이벤트 흐름 예시", level=1)
    bullets(doc, [
        "플레이어가 4발 선택 -> ASDBulletSelectorActor -> PlayerController.RequestRaiseTo(4).",
        "BettingSystem이 가능 여부 검사 -> GameState 베팅 상태 변경.",
        "GameState가 OnBetChanged.Broadcast(Player, 4) 발생.",
        "PresentationDirector가 이벤트 수신 -> 총알 4발 놓는 연출 실행.",
        "CRTMonitorActor와 버튼 상태 갱신.",
    ])

    doc.add_heading("10. 공통 이벤트 목록", level=1)
    bullets(doc, [
        "OnPhaseChanged(NewPhase)",
        "OnBetChanged(Side, BulletCount)",
        "OnCardsRevealed(PlayerCard, BossCard)",
        "OnRoundResolved(Result)",
        "OnRouletteStarted(Target, BulletCount)",
        "OnRouletteResult(Target, bHit)",
        "OnLifeChanged(Target, Life)",
        "OnStageChanged(Stage)",
        "OnGameOver(Winner)",
    ])

    doc.add_heading("11. 절대 주의", level=1)
    bullets(doc, [
        "코어 로직에서 연출, 사운드, 카메라를 직접 실행하지 않는다.",
        "연출 클래스에서 승패/확률/목숨 계산을 하지 않는다.",
        "LLM에게 보스 판단을 맡기지 않는다. LLM은 대사만 만든다.",
        "멀티에서 클라이언트가 결과를 직접 정하지 않는다. 서버/호스트가 결정한다.",
        "API Key, DB 비밀번호를 GitHub에 올리지 않는다.",
        "처음부터 모든 기능을 본 게임에 붙이지 말고 테스트맵에서 먼저 성공시킨다.",
    ])

    doc.add_heading("12. 한 줄 정리", level=1)
    doc.add_paragraph(
        "GameMode/GameState가 규칙을 돌리고, Interactable이 입력을 보내고, PresentationDirector가 이벤트를 받아 연출하며, LLM/DB/Multiplayer는 테스트맵에서 작게 성공시킨 뒤 연결한다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown C++ Class Structure")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

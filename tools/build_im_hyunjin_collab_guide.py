from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Role_ImHyunjin.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("임현진 역할 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("연출 총괄 / 사물형 UI 공동 작업 / 분위기 방향 / 레벨 통합 관리").bold = True

    doc.add_heading("1. 역할 목표", level=1)
    bullets(doc, [
        "현진님은 게임의 첫인상과 분위기를 잡는 역할입니다.",
        "단순한 카드 게임처럼 보이지 않게, 룰렛, 카드 공개, 베팅, 생명 카드 연출을 강하게 만듭니다.",
        "이진헌님과 함께 사물형 UI, 클릭 상호작용, 연출 시스템을 만듭니다.",
        "김윤아님이 찾고 배치한 에셋을 실제 게임 분위기에 맞게 조율합니다.",
        "메인 레벨/맵 통합을 관리해서 최종 화면이 한 방향으로 보이게 만듭니다.",
    ])

    doc.add_heading("2. 현진님이 먼저 만들 것", level=1)
    numbered(doc, [
        "ArtToneTest 맵: 테이블, 카드, 리볼버, 보스 실루엣, 조명으로 대표 분위기 만들기.",
        "EffectTest 맵: 버튼을 눌러 카드 공개, 룰렛, 생명 카드 소실 연출만 따로 테스트하기.",
        "메인 테이블 화면 구도: 플레이어가 가장 오래 볼 화면 잡기.",
        "기본 Post Process: 어두운 톤, 노이즈, 비네트, 낮은 채도 등 분위기 잡기.",
        "룰렛 연출: 총알 장전, 약실 회전, 방아쇠, 명중/불발.",
        "생명 카드 연출: 목숨이 줄어들 때 카드가 사라지거나 오염되는 효과.",
        "사물형 UI 방향: Call, Raise, Fold, 총알 선택, 체력 표시를 게임 안 사물로 표현.",
        "메인 레벨 통합 규칙: 누가 언제 레벨을 만지는지 정하고 충돌을 줄이기.",
    ])

    doc.add_heading("3. 이진헌님과 협업하는 방식", level=1)
    add_table(doc, ["현진님이 할 일", "이진헌님이 할 일"], [
        ["연출 방향과 화면 느낌 정하기.", "그 연출이 실제로 작동하도록 코드 연결."],
        ["버튼/스위치가 어떻게 눌려야 하는지 정하기.", "클릭/호버/비활성화 구조 만들기."],
        ["룰렛, 카드 공개, 생명 카드 연출 만들기.", "코어 이벤트가 오면 연출이 실행되게 연결."],
        ["메인 레벨 분위기와 최종 화면 관리.", "Git 병합과 기술 기능 통합 관리."],
        ["LLM/DB/멀티가 게임 분위기와 어울리게 보이도록 의견 주기.", "LLM/DB/멀티 테스트 기능 구현."],
    ])

    doc.add_heading("4. 박형빈님 코드와 연결하는 방식", level=1)
    bullets(doc, [
        "박형빈님 코드는 게임의 판단을 합니다.",
        "현진님은 그 판단 결과가 플레이어에게 잘 보이도록 연출을 만듭니다.",
        "예: 박형빈님 코드가 OnRouletteResult(Player, Hit) 이벤트 발생.",
        "이진헌님이 이벤트를 연출 쪽으로 연결.",
        "현진님/이진헌님이 명중 연출, 화면 암전, 생명 카드 소실을 실행.",
        "즉, 코어 판단 -> 이벤트 발생 -> 연출 실행 흐름으로 작업합니다.",
    ])

    doc.add_heading("5. 김윤아님과 협업하는 방식", level=1)
    add_table(doc, ["김윤아님이 할 일", "현진님이 할 일"], [
        ["분위기 레퍼런스와 에셋 후보 찾기.", "우리 게임 톤에 맞는 방향인지 같이 판단."],
        ["언리얼 테스트맵에 에셋 배치해보기.", "조명, 카메라, Post Process와 맞게 조율."],
        ["방 분위기 후보 2~3개 제안.", "최종 아트 톤과 연출 방향 선택."],
        ["카드, 총알, 버튼, 모니터 에셋 후보 찾기.", "사물형 UI와 연출에 쓸 수 있게 방향 잡기."],
        ["맵 드레싱 초안 만들기.", "메인 레벨에 반영할지 관리."],
    ])

    doc.add_heading("6. 레벨/맵 통합 관리", level=1)
    bullets(doc, [
        "메인 레벨 합치기와 최종 배치는 현진님이 관리합니다.",
        "각자 작업은 먼저 테스트맵에서 해보고, 메인 레벨 반영은 모여서 같이 확인합니다.",
        "메인 레벨을 수정하기 전에는 팀에 말합니다.",
        ".umap 파일은 충돌이 나면 해결이 어렵기 때문에 동시에 여러 명이 만지지 않습니다.",
        "윤아님이 배치한 에셋, 진헌님이 만든 상호작용, 현진님 연출이 한 화면에서 자연스럽게 보이는지 확인합니다.",
    ])

    doc.add_heading("7. 주의점", level=1)
    bullets(doc, [
        "연출에 집중하되, 코어 승패/확률/목숨 계산은 직접 하지 않습니다.",
        "어둡게 만드는 것과 안 보이게 만드는 것은 다릅니다. 카드, 총알, 버튼, 생명 카드는 잘 보여야 합니다.",
        "연출이 너무 길면 반복 플레이가 피곤해질 수 있습니다.",
        "사물형 UI는 멋있어야 하지만, 플레이어가 뭘 눌러야 하는지 모르면 안 됩니다.",
        "박형빈님 이벤트 이름이나 흐름을 바꿔야 할 때는 먼저 상의합니다.",
        "메인 레벨에 바로 작업하지 말고 테스트맵에서 먼저 확인합니다.",
    ])

    doc.add_heading("8. 완료 기준", level=1)
    bullets(doc, [
        "스크린샷만 봐도 게임 분위기가 느껴진다.",
        "카드 공개, 베팅, 룰렛, 목숨 감소가 연출로 잘 보인다.",
        "사물형 UI가 게임 안에 자연스럽게 녹아 있다.",
        "박형빈님 코어 이벤트가 오면 연출이 실행된다.",
        "진헌님과 만든 클릭/호버 상호작용이 실제 게임에서 작동한다.",
        "윤아님 에셋과 맵 배치가 최종 화면 안에서 어울린다.",
        "메인 레벨이 충돌 없이 관리되고, 팀이 같은 분위기를 보고 작업한다.",
    ])

    doc.add_heading("9. 한 줄 정리", level=1)
    doc.add_paragraph(
        "현진님은 게임의 분위기와 연출 방향을 잡는 담당입니다. 이진헌님과 함께 사물형 UI와 연출을 만들고, 박형빈님의 코어 이벤트를 화면으로 보여주며, 윤아님 에셋을 메인 레벨 안에서 하나의 분위기로 정리합니다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Im Hyunjin Role Guide")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

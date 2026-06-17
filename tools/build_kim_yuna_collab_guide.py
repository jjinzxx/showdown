from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Role_KimYuna.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("김윤아 역할 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("에셋 리서치 / 분위기 탐색 / 언리얼 배치 테스트 / 맵 드레싱").bold = True

    doc.add_heading("1. 역할 목표", level=1)
    bullets(doc, [
        "윤아님은 게임의 분위기와 공간 방향을 찾아가는 역할입니다.",
        "아직 최종 분위기와 스토리가 완전히 정해진 상태가 아니기 때문에, 에셋과 레퍼런스를 보면서 어떤 느낌이 좋을지 제안합니다.",
        "웹에서 에셋만 찾고 끝내는 것이 아니라, 언리얼 엔진에서 직접 배치해보며 실제 화면에서 어울리는지 확인합니다.",
        "최종 목표는 팀이 '이 분위기로 가자'라고 결정할 수 있게 화면 예시와 에셋 후보를 준비하는 것입니다.",
    ])

    doc.add_heading("2. 윤아님이 먼저 할 것", level=1)
    numbered(doc, [
        "분위기 레퍼런스 모으기: 어두운 방, 낡은 테이블, 아날로그 호러, 밀실, 가면, CRT, 조명.",
        "에셋 후보 찾기: 카드, 리볼버, 총알, 테이블, 의자, 조명, CRT, 버튼, 스위치, 후드 캐릭터, 마스크.",
        "언리얼 엔진에서 직접 배치해보기: 테스트맵에 넣고 실제 카메라 화면에서 확인.",
        "분위기 방향 2~3개 제안하기: 예를 들어 낡은 도박장, 아날로그 방송실, 수집가의 방.",
        "맵 드레싱 초안 만들기: 테이블 위에 카드, 총알, 리볼버, 버튼, 모니터를 어떻게 놓을지 배치.",
        "에셋 출처 정리하기: 어떤 에셋을 어디서 찾았는지 기록.",
        "스토리 단서 소품은 선택: 벽 낙서, 메모, 불탄 카드 등은 분위기에 맞으면 준비하고, 억지로 넣을 필요는 없음.",
    ])

    doc.add_heading("3. 찾아보면 좋은 분위기 방향", level=1)
    add_table(doc, ["방향", "느낌", "주의점"], [
        ["낡은 도박장 / 지하 게임장", "오래된 테이블, 어두운 조명, 낡은 카드, 녹슨 금속, 총알을 돈처럼 거는 느낌.", "너무 평범한 카지노처럼 보이면 호러가 약해질 수 있음."],
        ["아날로그 호러 / 낡은 방송실", "CRT 모니터, VHS 노이즈, 전선, 카메라, 녹음기, 이상한 방송 장치.", "소품이 많아지면 카드/총알이 안 보일 수 있음."],
        ["수집가의 방", "진열장, 이름이 적힌 카드, 마스크, 낡은 사진, 보관된 리볼버.", "장식이 너무 많으면 게임의 거친 느낌이 약해질 수 있음."],
    ])

    doc.add_heading("4. 에셋 우선순위", level=1)
    add_table(doc, ["우선순위", "찾을 것"], [
        ["1순위: 게임 플레이 필수", "카드, 리볼버, 총알/코인, 테이블, 버튼/스위치, CRT 모니터, 조명, 후드 보스, 마스크."],
        ["2순위: 분위기 소품", "벽 얼룩, 낡은 TV, 전선, 감시카메라, 진열장, 낡은 문, 녹슨 금속 소품, 오래된 가구."],
        ["3순위: 선택 사항", "벽 낙서, 찢어진 메모, 불탄 카드, 피 묻은 종이, 이전 플레이어 흔적, 스토리 문구."],
    ])

    doc.add_heading("5. 언리얼에서 배치할 때 볼 것", level=1)
    bullets(doc, [
        "실제 플레이 카메라에서 잘 보이는가?",
        "카드, 총알, 리볼버가 다른 소품에 묻히지 않는가?",
        "너무 어둡거나 너무 밝지는 않은가?",
        "테이블 위 핵심 사물이 한눈에 들어오는가?",
        "분위기는 좋은데 게임이 불편해지지는 않는가?",
        "임현진/이진헌이 사물형 UI로 쓰기 좋은 위치인가?",
        "에셋 크기가 서로 어색하지 않은가?",
    ])

    doc.add_heading("6. 임현진 / 이진헌과 협업하는 방식", level=1)
    add_table(doc, ["윤아님이 할 일", "임현진 / 이진헌이 할 일"], [
        ["레퍼런스와 에셋 후보 찾기.", "게임 분위기에 맞는지 같이 판단."],
        ["테스트맵에 직접 배치해보기.", "조명, 카메라, 연출과 맞게 조정."],
        ["분위기 방향 2~3개 제안.", "최종 아트 톤과 연출 방향 선택."],
        ["카드/총알/버튼 에셋 후보 찾기.", "클릭 상호작용과 연출에 맞게 사용."],
        ["맵 드레싱 초안 만들기.", "메인 레벨에 반영할지 결정."],
    ])

    doc.add_heading("7. 주의점", level=1)
    bullets(doc, [
        "웹에서 에셋만 찾고 끝내지 말기. 꼭 언리얼에 넣어서 확인하기.",
        "아직 분위기가 확정되지 않았으니 한 방향만 고집하지 않기.",
        "레퍼런스는 2~3개 방향으로 나눠서 모으기.",
        "너무 많은 소품을 넣어서 카드, 총알, 리볼버가 안 보이면 안 됨.",
        "분위기는 어두워도 플레이에 필요한 사물은 잘 보여야 함.",
        "에셋 출처와 용도를 꼭 기록하기.",
        "gta는 하루에 1시간만 하기.",
    ])

    doc.add_heading("8. 완료 기준", level=1)
    bullets(doc, [
        "분위기 후보 2~3개가 정리되어 있다.",
        "각 분위기 후보를 언리얼 테스트맵에서 간단히 배치해봤다.",
        "카드, 총, 총알, 테이블, 조명, CRT, 마스크 등 핵심 에셋 후보가 정리되어 있다.",
        "어떤 에셋을 어디서 찾았는지 출처가 정리되어 있다.",
        "팀이 최종 분위기를 고를 수 있을 만큼 화면 예시가 준비되어 있다.",
        "선택된 분위기에 맞춰 맵을 꾸밀 준비가 되어 있다.",
    ])

    doc.add_heading("9. 한 줄 정리", level=1)
    doc.add_paragraph(
        "윤아님은 에셋을 찾고, 언리얼 엔진에서 직접 배치해보면서 게임의 분위기 방향을 제안하는 역할입니다. 웹에서 자료만 모으는 것이 아니라, 실제 게임 화면에서 어울리는지 확인해서 팀이 최종 분위기를 고를 수 있게 도와줍니다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Kim Yuna Role Guide")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

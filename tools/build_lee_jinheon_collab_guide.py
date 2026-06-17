from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Role_LeeJinheon.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("이진헌 역할 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("기술 통합 / 연출 공동 작업 / 사물형 UI 공동 작업 / 멀티 / LLM / DB / Git 병합").bold = True

    doc.add_heading("1. 역할 목표", level=1)
    bullets(doc, [
        "진헌님은 기술적으로 어려운 기능들을 맡습니다.",
        "임현진과 함께 연출, 사물형 UI, 클릭 상호작용을 만듭니다.",
        "박형빈님이 만든 코어 게임 루프와 연출/UI가 서로 연결되게 만듭니다.",
        "LLM, DB, 멀티 기능을 작게라도 구현해서 평가 때 기술적으로 어필할 수 있게 합니다.",
        "Git 병합을 담당해서 팀 작업물이 develop 브랜치에 잘 합쳐지게 합니다.",
    ])

    doc.add_heading("2. 진헌님이 먼저 만들 것", level=1)
    numbered(doc, [
        "InteractableBase: 모든 클릭 가능한 사물의 기본 구조.",
        "DiegeticButton: 테이블 위 버튼이나 스위치. 예: Call, Raise, Fold.",
        "BulletSelector: Raise할 때 1~6발을 선택하는 장치.",
        "LLM_TestLevel: 보스 대사 응답을 받아보는 테스트맵.",
        "DB_TestLevel: 기록, 재화, 스킨 중 하나를 저장하고 불러오는 테스트맵.",
        "Multiplayer_TestMap: 1:1 접속과 상태 동기화를 테스트하는 맵.",
    ])

    doc.add_heading("3. 임현진과 협업하는 방식", level=1)
    add_table(doc, ["진헌님이 할 일", "임현진이 할 일"], [
        ["클릭 가능한 사물 구조 만들기.", "사물이 어떤 느낌으로 눌리고 보여야 하는지 정하기."],
        ["버튼/스위치가 코어 함수를 호출하게 만들기.", "버튼/스위치의 연출, 사운드, 분위기 잡기."],
        ["코어 이벤트와 연출 연결하기.", "이벤트가 왔을 때 어떤 연출이 나올지 만들기."],
        ["LLM/DB/멀티 기능 테스트하기.", "해당 기능이 게임 분위기와 어울리게 보이는 방식 잡기."],
        ["Git 병합 관리하기.", "레벨/맵 통합 관리하기."],
    ])

    doc.add_heading("4. 박형빈님 코드와 연결하는 방식", level=1)
    bullets(doc, [
        "박형빈님 코드는 게임의 판단을 합니다.",
        "진헌님은 그 판단 결과가 연출/UI로 이어지게 연결합니다.",
        "예: 박형빈님 코드가 OnBetChanged(Player, 4) 이벤트 발생.",
        "진헌님은 이 이벤트를 연출 쪽으로 전달.",
        "임현진/진헌님은 총알 4발이 테이블에 놓이는 연출 실행.",
        "즉, 코어 결과 발생 -> 진헌님이 연결 -> 임현진/진헌님 연출 실행.",
    ])

    doc.add_heading("5. LLM / DB / 멀티 작업 방식", level=1)
    doc.add_heading("LLM", level=2)
    bullets(doc, [
        "먼저 테스트맵에서 보스 대사 응답 받기.",
        "게임이 멈추지 않게 비동기로 처리하기.",
        "응답 실패 시 기본 대사 풀로 대체하기.",
        "LLM은 게임 판단을 하지 않고 대사만 담당.",
    ])
    doc.add_heading("DB", level=2)
    bullets(doc, [
        "처음부터 복잡하게 만들지 않기.",
        "기록, 재화, 스킨 중 하나만 먼저 저장/불러오기.",
        "게임을 껐다 켜도 데이터가 남아 있으면 성공.",
    ])
    doc.add_heading("멀티", level=2)
    bullets(doc, [
        "1:1 사설 방 또는 테스트 접속부터 시작.",
        "카드, 베팅, 공개, 룰렛 결과가 양쪽에서 같게 보이도록 하기.",
        "랭크, 시즌, 복잡한 매칭은 나중으로 미루기.",
    ])

    doc.add_heading("6. 주의점", level=1)
    bullets(doc, [
        "연출과 사물형 UI 방향은 임현진과 계속 상의하면서 맞추기.",
        "레벨/맵 통합은 임현진이 관리합니다.",
        "LLM, DB, 멀티는 본 게임에 바로 붙이지 말고 테스트맵에서 먼저 성공시키기.",
        "API Key, DB 비밀번호는 GitHub에 올리지 않기.",
        ".umap, .uasset 충돌은 혼자 억지로 해결하지 않기.",
        "멀티는 우선 1:1과 상태 동기화까지만 목표로 하기.",
        "Git 병합 전에는 에디터가 열리는지, 컴파일되는지 확인하기.",
    ])

    doc.add_heading("7. 완료 기준", level=1)
    bullets(doc, [
        "사물 클릭이 코어 함수로 연결된다.",
        "코어 이벤트가 연출/UI로 연결된다.",
        "임현진과 함께 만든 사물형 UI가 실제 게임에서 작동한다.",
        "LLM은 실패해도 기본 대사로 대체된다.",
        "DB 저장/불러오기가 된다.",
        "멀티 테스트에서 두 화면의 상태가 맞는다.",
        "develop 브랜치가 실행 가능한 상태로 유지된다.",
    ])

    doc.add_heading("8. 한 줄 정리", level=1)
    doc.add_paragraph(
        "진헌님은 기술 연결 담당입니다. 임현진과 함께 연출/UI를 만들고, 박형빈님의 코어 게임 루프를 연출과 연결하며, LLM/DB/멀티 기능을 작게 성공시켜 평가 때 기술력을 보여주는 역할입니다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Lee Jinheon Role Guide")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

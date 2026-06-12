from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Role_ParkHyungbin.docx"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 14, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": 80, "bottom": 80, "start": 120, "end": 120}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EEF5")
        margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    doc.add_paragraph("박형빈 역할 정리", style="Title")
    p = doc.add_paragraph()
    p.add_run("코어 게임 루프 / 룰 시스템 / 연출팀과 연결되는 이벤트 담당").bold = True

    doc.add_heading("1. 역할 목표", level=1)
    bullets(doc, [
        "형빈님은 게임이 규칙대로 돌아가게 만드는 코어 담당입니다.",
        "카드, 베팅, 승패, 룰렛, 목숨 계산을 정확하게 처리합니다.",
        "연출을 직접 만들기보다, 임현진/이진헌이 연출을 붙일 수 있도록 이벤트를 잘 알려줍니다.",
        "한마디로 형빈님 코드는 '무슨 일이 일어났는지'를 결정하고, 연출팀은 그 일을 '어떻게 보여줄지'를 맡습니다.",
    ])

    doc.add_heading("2. 형빈님이 먼저 만들 것", level=1)
    numbered(doc, [
        "임시 버튼이나 키보드 입력으로 한 판이 끝까지 돌아가게 만들기.",
        "덱 만들기: 1~7 각 2장, 총 14장.",
        "각자 손패 5장 지급, 4장 미사용 처리.",
        "카드 선택 후 상대 머리 위 카드로 등록.",
        "베팅 처리: Check, Raise, Call, Fold.",
        "카드 공개와 승패 판정.",
        "룰렛 결과 계산: 장전 수 / 6 확률.",
        "목숨 감소, 다음 판, 스테이지 클리어, 게임 오버 처리.",
        "간단한 보스 AI 만들기.",
    ])

    doc.add_heading("3. 연출팀과 협업하는 방식", level=1)
    add_table(doc, ["형빈님 코드가 할 일", "연출팀이 할 일"], [
        ["플레이어가 Raise 했는지 판단한다.", "총알이 테이블에 놓이는 연출을 만든다."],
        ["카드가 공개됐다는 이벤트를 보낸다.", "카드 뒤집기, 카메라 줌, 사운드를 재생한다."],
        ["누가 이겼는지 판정한다.", "승리/패배/동점 분위기를 연출한다."],
        ["룰렛 명중/불발을 계산한다.", "장전, 방아쇠, 암전, 생명 카드 소실을 보여준다."],
        ["목숨이 줄었다는 이벤트를 보낸다.", "생명 카드가 타거나 사라지는 효과를 만든다."],
    ])

    doc.add_heading("4. 꼭 만들어줘야 하는 이벤트", level=1)
    bullets(doc, [
        "OnPhaseChanged(NewPhase): 게임 단계가 바뀜.",
        "OnBetChanged(Side, BulletCount): 베팅 수가 바뀜.",
        "OnCardsRevealed(PlayerCard, BossCard): 카드가 공개됨.",
        "OnRoundResolved(Result): 승패/동점/폴드 결과가 결정됨.",
        "OnRouletteStarted(Target, BulletCount): 룰렛이 시작됨.",
        "OnRouletteResult(Target, bHit): 룰렛이 명중인지 불발인지 결정됨.",
        "OnLifeChanged(Target, Life): 목숨 수가 바뀜.",
        "OnStageChanged(Stage): 스테이지가 바뀜.",
        "OnGameOver(Winner): 게임이 끝남.",
    ])

    doc.add_heading("5. 함수는 이런 느낌이면 좋음", level=1)
    add_table(doc, ["함수", "역할"], [
        ["SelectCard(PlayerSide, CardId)", "선택 가능한 카드인지 확인하고 상대 머리 위 카드로 등록."],
        ["RaiseTo(PlayerSide, BulletCount)", "해당 총알 수로 Raise 가능한지 확인하고 베팅 반영."],
        ["Call(PlayerSide)", "상대 베팅 수에 맞춰 Call 처리."],
        ["Fold(PlayerSide)", "폴드 처리. 7 폴드 예외도 여기서 처리."],
        ["RevealCards()", "양쪽 카드를 공개하고 승패 판정으로 넘김."],
        ["ResolveRoulette(Target, BulletCount)", "Hit/Miss 계산 후 결과 이벤트 발생."],
        ["StartNextRound()", "카드 폐기 후 다음 판 준비."],
    ])

    doc.add_heading("6. 예시 흐름", level=1)
    bullets(doc, [
        "플레이어가 4발 Raise 버튼 클릭.",
        "이진헌/임현진의 사물형 UI가 RaiseTo(Player, 4)를 요청.",
        "형빈님 코드가 4발 Raise가 가능한지 검사.",
        "가능하면 베팅 상태를 바꾸고 OnBetChanged(Player, 4)를 발생.",
        "연출팀이 OnBetChanged를 받아 총알 4발 놓는 연출을 실행.",
    ])

    doc.add_heading("7. 주의점", level=1)
    bullets(doc, [
        "코어 코드 안에서 카메라, 사운드, VFX를 직접 실행하지 않기.",
        "코어는 계산과 이벤트 발생만 담당하기.",
        "연출팀이 필요한 정보는 이벤트로 알려주기.",
        "카드 상태를 명확히 나누기: 손패, 머리 위 카드, 버린 카드, 미사용 카드.",
        "처음부터 보스 AI를 어렵게 만들지 말고 Stage별 성향만 보이게 만들기.",
        "사물형 UI 완성을 기다리지 말고 임시 입력으로 코어 루프를 먼저 완성하기.",
        "버그 찾기 쉽게 현재 Phase, 카드 값, 베팅 수, 룰렛 결과를 로그로 남기기.",
    ])

    doc.add_heading("8. 완료 기준", level=1)
    bullets(doc, [
        "아트 없이도 카드 선택 -> 베팅 -> 공개 -> 룰렛 -> 목숨 감소 -> 다음 판이 돌아간다.",
        "7 폴드 예외, 동점 양쪽 룰렛, 6발 확정 명중이 동작한다.",
        "임현진/이진헌이 받을 이벤트가 모두 발생한다.",
        "연출이 없어도 게임 규칙이 맞게 돌아간다.",
        "연출을 붙이면 바로 화면에 반영될 정도로 이벤트 연결이 깔끔하다.",
    ])

    doc.add_heading("9. 한 줄 정리", level=1)
    doc.add_paragraph(
        "형빈님은 게임의 판단 담당입니다. 연출을 직접 만들지 말고, 카드/베팅/룰렛/목숨 결과를 정확히 계산한 뒤 이벤트로 알려주면 임현진과 이진헌이 그 이벤트에 맞춰 연출을 붙입니다."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Park Hyungbin Role Guide")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

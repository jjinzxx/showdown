from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 15, "1F4D78"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)


def add_list(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        doc.add_paragraph(item, style=style)


def add_section(doc, title, items, numbered=False):
    doc.add_heading(title, level=1)
    add_list(doc, items, numbered=numbered)


def make_doc(title, subtitle, sections, filename):
    doc = Document()
    setup(doc)
    doc.add_paragraph(title, style="Title")
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.bold = True

    for section_title, items, numbered in sections:
        add_section(doc, section_title, items, numbered)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Simple Guide")

    DOCS.mkdir(exist_ok=True)
    path = DOCS / filename
    doc.save(path)
    return path


def main():
    docs = [
        (
            "임현진 역할 정리",
            "연출 메인 / 사물형 UI / VFX",
            [
                ("해야 할 것", [
                    "메인 테이블 화면 분위기 잡기: 조명, 카메라, Post Process.",
                    "사물형 UI 방향 잡기: 체력, 베팅, Call/Raise/Fold를 게임 안 사물로 표현.",
                    "룰렛 연출 만들기: 장전, 약실 회전, 방아쇠, 명중/불발.",
                    "생명 카드 소실 연출 만들기.",
                    "카드 공개, 베팅, 폴드 순간의 화면/사운드 피드백 만들기.",
                    "Stage가 올라갈수록 방, 보스, 화면 노이즈가 강해지는 느낌 잡기.",
                ], False),
                ("먼저 만들 것", [
                    "LV_ArtToneTest: 분위기 테스트용 맵.",
                    "LV_EffectTest: 버튼으로 연출만 테스트하는 맵.",
                    "BP_RouletteSequence: 룰렛 연출 묶음.",
                    "BP_LifeCardDisplay: 목숨 표시/소실 연출.",
                ], False),
                ("주의점", [
                    "코어 로직을 직접 계산하지 말기. 승패, 확률, 목숨 계산은 박형빈 담당.",
                    "어둡게 만들되 카드/총알/버튼은 반드시 잘 보여야 함.",
                    "모든 UI를 사물형으로 만들려고 욕심내지 말고 핵심 플레이 UI부터.",
                    "연출은 너무 길면 반복 플레이가 피곤해짐.",
                ], False),
                ("완료 기준", [
                    "스크린샷만 봐도 게임 분위기가 느껴진다.",
                    "룰렛 장면이 평가 때 가장 기억에 남는다.",
                    "코어 이벤트가 들어오면 연출이 안정적으로 재생된다.",
                ], False),
            ],
            "ShowDown_Role_ImHyunjin.docx",
        ),
        (
            "이진헌 역할 정리",
            "기술 통합 / 연출 보조 / 멀티 / LLM / DB / Git",
            [
                ("해야 할 것", [
                    "Git 병합 담당: develop 브랜치 관리.",
                    "사물형 UI 클릭/호버 시스템 만들기.",
                    "임현진의 연출과 박형빈의 코어 이벤트 연결.",
                    "LLM 테스트: 보스 대사 응답 받기.",
                    "DB 테스트: 기록/재화/스킨 중 하나 저장하고 불러오기.",
                    "멀티 테스트: 1:1 접속과 핵심 상태 동기화.",
                ], False),
                ("먼저 만들 것", [
                    "BP_InteractableBase: 모든 클릭 사물의 공통 부모.",
                    "BP_DiegeticButton: 테이블 버튼/스위치.",
                    "BP_BulletSelector: 1~6발 선택 장치.",
                    "LLM_TestLevel, DB_TestLevel, Multiplayer_TestMap.",
                ], False),
                ("주의점", [
                    "LLM, DB, 멀티는 본 게임에 바로 붙이지 말고 테스트맵에서 먼저 성공시키기.",
                    "API Key, DB 비밀번호는 GitHub에 올리지 않기.",
                    "충돌 난 .umap/.uasset은 혼자 억지로 해결하지 않기.",
                    "멀티는 사설 1:1과 상태 동기화까지만 우선.",
                ], False),
                ("완료 기준", [
                    "사물 클릭이 코어 함수로 연결된다.",
                    "LLM은 실패해도 대사 풀로 대체된다.",
                    "DB 저장/불러오기가 된다.",
                    "멀티 테스트에서 두 화면의 상태가 맞는다.",
                ], False),
            ],
            "ShowDown_Role_LeeJinheon.docx",
        ),
        (
            "박형빈 역할 정리",
            "코어 게임 루프 / 룰 시스템 / 기본 보스 AI",
            [
                ("해야 할 것", [
                    "덱 만들기: 1~7 각 2장, 총 14장.",
                    "각자 손패 5장 지급, 4장 미사용 처리.",
                    "카드 선택 → 상대 머리 위 카드 등록.",
                    "베팅 구현: Check, Raise, Call, Fold.",
                    "카드 공개와 승패 판정.",
                    "룰렛 확률 계산: 장전 수 / 6.",
                    "목숨 감소, 다음 판, 스테이지 클리어, 게임 오버 처리.",
                    "간단한 보스 AI 만들기.",
                ], False),
                ("먼저 만들 것", [
                    "임시 2D 버튼이나 키보드 입력으로 한 판이 끝까지 돌게 만들기.",
                    "GameManager, CardSystem, BettingSystem, RouletteSystem.",
                    "OnBetChanged, OnCardsRevealed, OnRouletteResult 같은 이벤트 발생.",
                ], False),
                ("주의점", [
                    "연출/VFX/사운드를 코어 안에서 직접 실행하지 않기.",
                    "코어는 규칙 계산과 이벤트 발생만 담당.",
                    "카드 상태를 명확히 나누기: 손패, 머리 위 카드, 버린 카드, 미사용 카드.",
                    "AI는 처음부터 복잡하게 만들지 말고 Stage별 성향만 보이게.",
                ], False),
                ("완료 기준", [
                    "아트 없이도 게임 한 판이 끝까지 돌아간다.",
                    "7 폴드 예외, 동점 양쪽 룰렛, 6발 확정 명중이 동작한다.",
                    "연출 담당자가 받을 이벤트가 모두 발생한다.",
                ], False),
            ],
            "ShowDown_Role_ParkHyungbin.docx",
        ),
        (
            "김윤아 역할 정리",
            "에셋 / 맵 드레싱 / 스토리 소품",
            [
                ("해야 할 것", [
                    "필요 에셋 찾기: 카드, 리볼버, 총알, 테이블, CRT, 조명, 후드 캐릭터.",
                    "에셋 출처와 라이선스 정리.",
                    "방 블록아웃 만들기: 플레이어 자리, 보스 자리, 테이블, 카메라 위치.",
                    "테이블 위 핵심 소품 배치.",
                    "벽 낙서, 메모, 타다 만 카드 같은 스토리 소품 준비.",
                    "Stage별 방 변화 아이디어 정리.",
                ], False),
                ("먼저 만들 것", [
                    "에셋 리스트 문서.",
                    "방 레퍼런스 이미지 모음.",
                    "LV_Environment_Blockout.",
                    "핵심 소품 1차 배치.",
                ], False),
                ("주의점", [
                    "무료 에셋도 라이선스 확인하기.",
                    "메인 맵은 동시에 여러 명이 수정하지 않기.",
                    "소품을 너무 많이 넣어서 카드/총알이 안 보이게 만들지 않기.",
                    "실제 유명인 얼굴을 그대로 쓰지 말고 변형된 오리지널 마스크 방향으로 잡기.",
                ], False),
                ("완료 기준", [
                    "메인 방이 빈 맵처럼 보이지 않는다.",
                    "핵심 소품이 준비되어 있고 출처가 기록되어 있다.",
                    "Stage별 분위기 변화가 눈에 보인다.",
                ], False),
            ],
            "ShowDown_Role_KimYuna.docx",
        ),
        (
            "공통 협업 규칙",
            "팀 전체가 지켜야 할 최소 규칙",
            [
                ("역할 요약", [
                    "임현진: 연출 메인, 사물형 UI, VFX.",
                    "이진헌: 기술 통합, 연출 보조, 멀티, LLM, DB, Git 병합.",
                    "박형빈: 코어 게임 루프, 룰 시스템, 기본 보스 AI.",
                    "김윤아: 에셋, 맵 드레싱, 스토리 소품.",
                ], False),
                ("협업 방식", [
                    "박형빈은 코어 이벤트를 만든다.",
                    "임현진은 이벤트를 받아 연출을 재생한다.",
                    "이진헌은 코어와 연출, 사물형 UI를 연결한다.",
                    "김윤아는 맵과 소품을 준비한다.",
                    "각자 테스트맵에서 먼저 만들고, 통합은 develop에서 한다.",
                ], False),
                ("Git 규칙", [
                    "main: 안정 버전.",
                    "develop: 통합 버전.",
                    "feature/core-loop: 박형빈.",
                    "feature/presentation: 임현진.",
                    "feature/tech-integration: 이진헌.",
                    "feature/environment: 김윤아.",
                    "커밋은 작게, 메시지는 알아보기 쉽게.",
                ], False),
                ("절대 주의", [
                    ".umap, .uasset은 동시에 수정하지 않기.",
                    "Binaries, Intermediate, Saved, DerivedDataCache는 Git에 올리지 않기.",
                    "API Key, DB 비밀번호는 GitHub에 올리지 않기.",
                    "하루 이상 막히면 바로 공유하기.",
                    "평가 직전에는 새 기능 추가보다 안정화 우선.",
                ], False),
                ("우선순위", [
                    "1순위: 싱글 코어 루프.",
                    "2순위: 사물형 UI와 룰렛 연출.",
                    "3순위: LLM/DB/멀티 테스트 성공.",
                    "4순위: 전체 통합과 시연 안정화.",
                ], False),
            ],
            "ShowDown_Collaboration_Common.docx",
        ),
    ]

    for args in docs:
        print(make_doc(*args))


if __name__ == "__main__":
    main()

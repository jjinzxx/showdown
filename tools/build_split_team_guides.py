from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_team_collaboration_guide import (
    add_bullets,
    add_callout,
    add_matrix,
    add_numbered,
    setup_styles,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def save_doc(doc, filename):
    DOCS.mkdir(exist_ok=True)
    path = DOCS / filename
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Split Team Guide")
    doc.save(path)
    return path


def base_doc(title, subtitle):
    doc = Document()
    setup_styles(doc)
    doc.add_paragraph("ShowDown", style="Title")
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    doc.add_paragraph(subtitle)
    return doc


def build_im():
    doc = base_doc(
        "임현진 역할 가이드: 연출 메인 / 사물형 UI / VFX",
        "목표: 단순한 카드 게임 루프를 평가자가 바로 기억할 수 있는 장면, 사운드, 사물형 UI, 룰렛 연출로 끌어올린다.",
    )
    add_callout(doc, "임현진의 핵심 책임", [
        "플레이어가 보는 첫 화면, 카드 공개, 베팅, 룰렛, 생명 카드 소실이 강하게 느껴지도록 만든다.",
        "코어 규칙 계산을 직접 하지 않고, 코어 이벤트를 받아 연출을 재생하는 구조로 작업한다.",
        "2D UI를 최소화하고 체력, 베팅, 선택, 위험도를 게임 안 사물로 표현한다.",
    ])

    doc.add_heading("1. 가장 먼저 만들 것", level=1)
    add_numbered(doc, [
        "LV_ArtToneTest: 어두운 방, 테이블, 보스 실루엣, 카드, 총, 총알, 생명 카드를 배치한 대표 장면.",
        "LV_EffectTest: 버튼만 눌러 카드 공개, 총알 장전, 격발, 명중/불발, 생명 카드 소실을 각각 테스트하는 맵.",
        "BP_PresentationDirector: 코어 이벤트를 받아 연출을 재생하는 중앙 관리자.",
        "BP_RouletteSequence: BulletCount와 Hit/Miss 값을 받아 룰렛 연출을 재생하는 블루프린트.",
        "BP_LifeCardDisplay: 목숨 수에 따라 테이블 위 생명 카드가 사라지거나 오염되는 연출.",
        "PP_AnalogHorror_Base: 낮은 채도, 높은 대비, 비네트, 필름 그레인, 약한 색수차가 들어간 기본 포스트프로세스.",
    ])

    doc.add_heading("2. 담당 연출 목록", level=1)
    add_bullets(doc, [
        "메인 테이블 화면 구도와 카메라 위치.",
        "카드 선택 및 카드 공개 연출.",
        "Call, Raise, Fold 선택 시 사물형 버튼 피드백.",
        "Raise 1~6발 선택 장치의 시각/사운드 피드백.",
        "총알이 테이블에 놓이거나 리볼버에 장전되는 연출.",
        "방아쇠 당기기 전 정적, 심장박동, 카메라 긴장감.",
        "명중/불발 결과 연출.",
        "생명 카드 소실, 오염, 불타는 효과.",
        "Stage별 조명, 노이즈, 보스 실루엣 변화.",
        "보스 마스크가 조명에 따라 불쾌하게 보이는 연출.",
    ])

    doc.add_heading("3. 블루프린트 작성 방식", level=1)
    add_bullets(doc, [
        "연출은 함수처럼 호출 가능하게 만든다. 예: PlayRevealCards, PlayRouletteSequence, PlayLifeLost.",
        "처음에는 테스트 버튼으로 실행되게 만들고, 나중에 코어 이벤트와 연결한다.",
        "연출 중 입력 잠금이 필요한 경우 BP_PresentationDirector에서 관리한다.",
        "카메라, 사운드, 조명, VFX 타이밍은 Sequence 단위로 묶는다.",
        "Stage별 노이즈 강도, 조명 깜빡임, 카메라 흔들림은 변수나 DataAsset으로 빼서 조정 가능하게 둔다.",
    ])

    doc.add_heading("4. 주의점", level=1)
    add_bullets(doc, [
        "어둡게 만드는 것과 안 보이게 만드는 것은 다르다. 카드, 총알, 리볼버, 생명 카드는 항상 읽혀야 한다.",
        "모든 메뉴를 사물형 UI로 만들려고 하지 말고, 플레이 중 핵심 UI부터 완성한다.",
        "연출이 반복될수록 피곤해질 수 있으니 첫 연출과 반복 연출의 길이를 다르게 둘 수 있게 만든다.",
        "코어 로직 안에 직접 들어가서 승패나 확률을 계산하지 않는다.",
        "메인 맵을 직접 오래 만지기보다 테스트 맵에서 먼저 완성하고 통합 시점에 붙인다.",
    ])

    doc.add_heading("5. 완료 기준", level=1)
    add_bullets(doc, [
        "스크린샷 한 장만 봐도 죽음의 카드 베팅 게임처럼 보인다.",
        "카드 공개, 베팅, 룰렛, 목숨 감소가 2D 설명 없이도 이해된다.",
        "룰렛 장면이 평가 시연에서 가장 기억에 남는 장면이 된다.",
        "코어 이벤트가 들어오면 연출이 안정적으로 실행된다.",
    ])
    return save_doc(doc, "ShowDown_Role_ImHyunjin.docx")


def build_lee():
    doc = base_doc(
        "이진헌 역할 가이드: 기술 통합 / 연출 보조 / 멀티 / LLM / DB / Git",
        "목표: 팀원이 만든 코어, 연출, 에셋을 안정적으로 연결하고, 평가에서 보여줄 기술 쇼케이스 기능을 작게 성공시킨다.",
    )
    add_callout(doc, "이진헌의 핵심 책임", [
        "Git 병합과 통합 브랜치를 관리한다.",
        "사물형 UI가 코어 게임 함수를 호출할 수 있도록 인터랙션 시스템을 만든다.",
        "멀티, LLM, DB는 본 게임에 바로 붙이지 말고 테스트 레벨에서 먼저 성공시킨다.",
    ])

    doc.add_heading("1. 가장 먼저 만들 것", level=1)
    add_numbered(doc, [
        "Git 브랜치 구조 정리: main, develop, feature/core-loop, feature/presentation, feature/tech-integration, feature/environment.",
        "BP_InteractableBase: Hover, Unhover, Interact, Enable, Disable을 가진 공통 상호작용 부모.",
        "BP_DiegeticButton: 테이블 버튼, 레버, 스위치 같은 사물형 버튼의 공통 구현.",
        "BP_BulletSelector: 1~6발 선택 후 Core의 RaiseTo(BulletCount)를 호출하는 장치.",
        "LLM_TestLevel: 본 게임과 분리된 API 응답 테스트.",
        "DB_TestWidget 또는 DB_TestLevel: 기억 조각, 승패 기록, 스킨 해금 중 하나를 저장/불러오기.",
        "Multiplayer_TestMap: 두 클라이언트 접속과 간단한 상태 동기화 검증.",
    ])

    doc.add_heading("2. 담당 기능 목록", level=1)
    add_bullets(doc, [
        "Git 병합, 충돌 관리, develop 브랜치 안정화.",
        "사물 클릭/호버/비활성화 공통 시스템.",
        "사물형 버튼과 코어 함수 연결.",
        "코어 이벤트와 임현진 연출 블루프린트 연결.",
        "LLM 호출, 비동기 처리, 실패 시 대사 풀 폴백.",
        "DB 저장/불러오기 기능.",
        "멀티 1:1 테스트와 핵심 상태 동기화.",
    ])

    doc.add_heading("3. 코드 작성 방식", level=1)
    add_bullets(doc, [
        "Interactable은 게임 규칙을 계산하지 않는다. 클릭되면 GameManager의 공개 함수만 호출한다.",
        "LLM은 보스 행동을 결정하지 않고 대사만 만든다. 판단은 코드가 한다.",
        "LLM 응답이 늦으면 게임이 멈추지 않게 하고, 대사 풀로 즉시 대체한다.",
        "DB는 처음부터 계정/보안/랭크를 다 만들지 말고 평가 시연에 필요한 저장 하나를 확실히 만든다.",
        "멀티에서는 카드 지급, 승패, 룰렛 결과를 서버/호스트 권위로 처리한다.",
        "통합 전에는 에디터 실행, 컴파일, 주요 테스트맵 실행을 확인한다.",
    ])

    doc.add_heading("4. 주의점", level=1)
    add_bullets(doc, [
        "API Key, DB 비밀번호는 절대 GitHub에 올리지 않는다.",
        "멀티, LLM, DB를 동시에 본 게임에 붙이면 디버깅이 어려우므로 반드시 분리 테스트한다.",
        "충돌이 난 .umap, .uasset은 혼자 억지로 해결하지 말고 해당 파일 담당자와 같이 본다.",
        "통합 담당이라고 모든 버그를 혼자 고치지 않는다. 담당자에게 재현 방법과 로그를 전달한다.",
        "멀티 범위는 1:1 사설 방과 핵심 상태 동기화에 집중한다.",
    ])

    doc.add_heading("5. 완료 기준", level=1)
    add_bullets(doc, [
        "사물형 UI 클릭이 코어 함수로 안정적으로 연결된다.",
        "코어 이벤트가 연출 블루프린트를 호출한다.",
        "LLM은 응답 성공과 실패 폴백이 모두 동작한다.",
        "DB는 저장 후 재실행해도 데이터가 다시 불러와진다.",
        "멀티 테스트에서 두 클라이언트의 핵심 상태가 맞게 보인다.",
        "develop 브랜치가 시연 가능한 상태로 유지된다.",
    ])
    return save_doc(doc, "ShowDown_Role_LeeJinheon.docx")


def build_park():
    doc = base_doc(
        "박형빈 역할 가이드: 코어 게임 루프 / 룰 시스템 / 기본 보스 AI",
        "목표: 아트와 연출이 없어도 카드 선택부터 룰렛 결과까지 한 판이 끝까지 안정적으로 돌아가는 싱글 코어를 만든다.",
    )
    add_callout(doc, "박형빈의 핵심 책임", [
        "게임 규칙, 상태, 승패, 룰렛 확률을 정확하게 처리한다.",
        "연출을 직접 실행하지 않고 상태 변화 이벤트만 발생시킨다.",
        "처음에는 임시 2D 버튼이나 키보드 입력으로라도 코어 루프를 빠르게 완성한다.",
    ])

    doc.add_heading("1. 가장 먼저 만들 것", level=1)
    add_numbered(doc, [
        "BP_GameManager 또는 C++ GameManager: Phase, 턴, 목숨, 스테이지 상태 관리.",
        "CardSystem: 1~7 각 2장, 총 14장 덱 생성, 각자 5장 지급, 4장 미사용 처리.",
        "BettingSystem: Check, RaiseTo, Call, Fold 처리. 한 판 최대 6발 제한.",
        "RoundResolver: 카드 공개, 높은 숫자 승리, 동점 양쪽 룰렛 처리.",
        "RouletteSystem: BulletCount / 6 확률로 Hit/Miss 결정. 6발은 100% 명중.",
        "LifeSystem: 목숨 감소, 사망, 스테이지 클리어, 게임 오버 처리.",
        "BossAI_Basic: Stage별로 보수적/공격적 베팅 성향을 바꾸는 간단 AI.",
    ])

    doc.add_heading("2. 공개 함수 목록", level=1)
    add_bullets(doc, [
        "SelectCard(PlayerId, CardId): 선택 가능한 카드인지 검사하고 상대 머리 위 카드로 등록.",
        "RaiseTo(PlayerId, BulletCount): 현재 베팅보다 큰지, 1~6 범위인지 검사.",
        "Call(PlayerId): 상대 최종 베팅 수와 맞춘다.",
        "Fold(PlayerId): 그 판 패배 처리. 7 폴드 예외면 6발 강제.",
        "RevealCards(): 양쪽 카드 공개 후 승패 판정.",
        "ResolveRoulette(TargetPlayer, BulletCount): 명중 여부 계산 후 결과 이벤트 발생.",
        "StartNextRound(): 카드 폐기, 다음 카드 선택 단계로 전환.",
    ])

    doc.add_heading("3. 발생시켜야 할 이벤트", level=1)
    add_bullets(doc, [
        "OnPhaseChanged(NewPhase)",
        "OnCardSelected(Player, CardId)",
        "OnBetChanged(Player, BulletCount)",
        "OnActionResolved(ActionType)",
        "OnCardsRevealed(PlayerCard, BossCard)",
        "OnRoundResolved(Result)",
        "OnRouletteStarted(Target, BulletCount)",
        "OnRouletteResult(Target, bHit)",
        "OnLifeChanged(Target, Life)",
        "OnStageChanged(Stage)",
        "OnGameOver(Winner)",
    ])

    doc.add_heading("4. 주의점", level=1)
    add_bullets(doc, [
        "코어 로직 안에서 카메라 흔들림, 사운드, VFX를 직접 재생하지 않는다.",
        "카드 상태를 명확히 나눈다: 손패, 머리 위 카드, 공개 카드, 버린 카드, 미사용 카드.",
        "한 함수가 너무 많은 일을 하지 않게 한다. Fold는 폴드 처리, ResolveRoulette는 룰렛 결과 계산처럼 책임을 나눈다.",
        "디버그 로그를 적극적으로 남긴다. 현재 Phase, 카드 값, 베팅 수, 룰렛 결과 로그가 중요하다.",
        "AI는 처음부터 똑똑하게 만들지 말고 Stage별 성향이 보이는 단순 AI부터 만든다.",
        "사물형 UI 완성을 기다리지 말고 임시 UI로 코어 루프를 먼저 돌린다.",
    ])

    doc.add_heading("5. 완료 기준", level=1)
    add_bullets(doc, [
        "임시 UI만으로도 카드 선택, 베팅, 공개, 룰렛, 목숨 감소, 다음 판이 반복된다.",
        "7 폴드 예외, 동점 양쪽 룰렛, 6발 확정 명중이 테스트된다.",
        "목숨 0, 스테이지 클리어, 게임 오버가 정상 처리된다.",
        "연출/상호작용 담당자가 받을 이벤트가 모두 발생한다.",
    ])
    return save_doc(doc, "ShowDown_Role_ParkHyungbin.docx")


def build_kim():
    doc = base_doc(
        "김윤아 역할 가이드: 에셋 / 맵 드레싱 / 스토리 소품",
        "목표: 게임 공간, 소품, 방 분위기, 스토리 단서를 준비해서 단순한 테이블 게임이 아니라 하나의 불쾌한 방처럼 느껴지게 만든다.",
    )
    add_callout(doc, "김윤아의 핵심 책임", [
        "에셋을 찾는 것에서 끝나지 않고, 실제 맵에 어울리게 배치하고 정리한다.",
        "카드, 리볼버, 총알, 생명 카드, 테이블처럼 게임 정보와 연결되는 소품을 우선한다.",
        "에셋 출처와 라이선스를 기록해서 평가/배포 리스크를 줄인다.",
    ])

    doc.add_heading("1. 가장 먼저 만들 것", level=1)
    add_numbered(doc, [
        "에셋 리스트 스프레드시트: 이름, 용도, 출처, 라이선스, 다운로드 링크, 사용 위치.",
        "방 콘셉트 보드: 좁은 방, 낡은 테이블, 아날로그 호러, CRT, 후드 보스 레퍼런스.",
        "LV_Environment_Blockout: 테이블, 보스 자리, 플레이어 자리, 조명, 카메라 위치가 잡힌 맵 초안.",
        "핵심 소품 세트: 카드, 리볼버, 총알/코인, 생명 카드, 테이블, 조명, CRT.",
        "스토리 단서 소품: 벽 긁힌 글자, 타다 만 카드, 메모, 이전 희생자의 흔적.",
        "Stage별 변화안: Stage 1은 정돈, Stage 2는 오염, Stage 3은 붕괴와 왜곡.",
    ])

    doc.add_heading("2. 담당 작업 목록", level=1)
    add_bullets(doc, [
        "무료/사용 가능 에셋 리서치와 정리.",
        "맵 블록아웃과 방 크기 조정.",
        "테이블 위 핵심 소품 배치.",
        "CRT, 낡은 조명, 벽 낙서, 종이 메모 같은 분위기 소품 배치.",
        "보스 후드/마스크 레퍼런스 수집.",
        "Stage별 방 변화용 오브젝트 정리.",
        "스토리 단서 텍스트와 오브젝트 초안 준비.",
    ])

    doc.add_heading("3. 작업 방식", level=1)
    add_bullets(doc, [
        "에셋 파일명은 역할이 보이게 짓는다. 예: SM_Table_Rusty, M_Card_Burnt, BP_CRT_Monitor.",
        "폴더를 정리한다. 예: Content/ShowDown/Environment, Props, Characters, Materials, UIObjects, VFX.",
        "에셋은 실제 플레이 카메라에서 확인한다. 에디터 뷰에서 좋아도 플레이 화면에서 안 보이면 수정한다.",
        "소품은 플레이 가독성을 해치지 않게 배치한다. 카드, 총알, 리볼버가 안 보이면 안 된다.",
        "보스 마스크는 실제 유명인 얼굴을 그대로 쓰지 말고, 변형된 오리지널 마스크 방향으로 잡는다.",
    ])

    doc.add_heading("4. 주의점", level=1)
    add_bullets(doc, [
        "무료 에셋이라도 라이선스가 다르다. 출처 표기 필요 여부를 꼭 기록한다.",
        "메인 맵 파일은 동시에 여러 명이 만지지 않는다. 작업 전후에 팀에게 알린다.",
        "너무 많은 고해상도 에셋을 넣으면 프로젝트가 무거워진다. 화면에 보이는 핵심부터 넣는다.",
        "전체 방을 예쁘게 채우는 것보다 게임의 핵심 사물과 분위기가 먼저다.",
        "임현진의 Post Process와 조명 아래에서 색감이 어떻게 보이는지 같이 확인한다.",
    ])

    doc.add_heading("5. 완료 기준", level=1)
    add_bullets(doc, [
        "메인 테이블 장면이 빈 회색박스가 아니라 하나의 방처럼 보인다.",
        "핵심 소품이 모두 준비되어 있고 출처가 기록되어 있다.",
        "Stage별 방 변화 방향이 시각적으로 구분된다.",
        "스토리 단서 오브젝트가 맵 안에 자연스럽게 들어갈 준비가 되어 있다.",
    ])
    return save_doc(doc, "ShowDown_Role_KimYuna.docx")


def build_collab():
    doc = base_doc(
        "공통 협업/주의사항 가이드",
        "목표: 초보 4인 팀이 Unreal 프로젝트에서 충돌을 줄이고, 코어 로직과 연출을 안정적으로 연결하기 위한 공통 규칙.",
    )
    add_callout(doc, "공통 원칙", [
        "코어는 규칙만 처리하고, 연출은 이벤트를 받아 반응한다.",
        "같은 .umap, .uasset을 동시에 만지지 않는다.",
        "멀티, LLM, DB는 테스트 레벨에서 작게 성공시킨 뒤 본 게임에 붙인다.",
    ])

    doc.add_heading("1. 역할 요약", level=1)
    add_matrix(
        doc,
        ["팀원", "주 담당", "핵심 산출물"],
        [
            ["임현진", "연출 메인 / 사물형 UI / VFX", "ArtToneTest, EffectTest, PresentationDirector, RouletteSequence"],
            ["이진헌", "기술 통합 / 연출 보조 / 멀티 / LLM / DB / Git", "InteractableBase, BulletSelector, LLM/DB/Multiplayer 테스트, develop 통합"],
            ["박형빈", "코어 게임 루프 / 룰 시스템 / 기본 보스 AI", "GameManager, CardSystem, BettingSystem, RouletteSystem, LifeSystem"],
            ["김윤아", "에셋 / 맵 드레싱 / 스토리 소품", "에셋 리스트, Environment Blockout, 핵심 소품, Stage별 방 변화"],
        ],
        [1500, 3000, 4860],
    )

    doc.add_heading("2. 코드 구조", level=1)
    add_matrix(
        doc,
        ["층", "책임", "예시"],
        [
            ["Core", "게임 규칙과 상태 처리", "덱, 손패, 베팅, 승패, 룰렛, 목숨, 스테이지"],
            ["Interaction", "사물 클릭/조작 입력 처리", "테이블 버튼, 총알 선택 장치, 카드 선택 오브젝트"],
            ["Presentation", "카메라, 사운드, VFX, UI 피드백", "룰렛 시퀀스, 생명 카드 소실, 화면 노이즈"],
            ["Content", "맵, 소품, 텍스트, 에셋 배치", "방 드레싱, CRT, 벽 낙서, 보스 마스크"],
        ],
        [1500, 2500, 5360],
    )

    doc.add_heading("3. 공통 이벤트", level=1)
    add_bullets(doc, [
        "OnPhaseChanged(NewPhase)",
        "OnCardSelected(Player, CardId)",
        "OnBetChanged(Player, BulletCount)",
        "OnActionResolved(ActionType)",
        "OnCardsRevealed(PlayerCard, BossCard)",
        "OnRoundResolved(Result)",
        "OnRouletteStarted(Target, BulletCount)",
        "OnRouletteResult(Target, bHit)",
        "OnLifeChanged(Target, Life)",
        "OnStageChanged(Stage)",
        "OnGameOver(Winner)",
    ])

    doc.add_heading("4. Git 규칙", level=1)
    add_matrix(
        doc,
        ["브랜치", "담당", "용도"],
        [
            ["main", "이진헌", "시연 가능한 안정 버전"],
            ["develop", "이진헌", "통합 브랜치"],
            ["feature/core-loop", "박형빈", "싱글 코어 루프"],
            ["feature/presentation", "임현진", "연출, VFX, 사물형 UI 콘셉트"],
            ["feature/tech-integration", "이진헌", "인터랙션, LLM, DB, 멀티"],
            ["feature/environment", "김윤아", "맵, 에셋, 소품 배치"],
        ],
        [2300, 1600, 5460],
    )
    add_bullets(doc, [
        "커밋은 작게 한다.",
        "커밋 전에 에디터가 열리는지 확인한다.",
        "공용 파일을 수정해야 하면 먼저 단체 채팅에 말한다.",
        "충돌이 나면 담당자 없이 억지로 해결하지 않는다.",
        "Binaries, Intermediate, Saved, DerivedDataCache는 Git에 올리지 않는다.",
        "API Key, DB 비밀번호, 개인 설정 파일은 절대 커밋하지 않는다.",
    ])

    doc.add_heading("5. Unreal 협업 주의점", level=1)
    add_bullets(doc, [
        ".umap과 .uasset은 Git에서 충돌 해결이 어렵다. 파일 소유자를 정한다.",
        "메인 레벨은 한 명만 수정하는 시간을 정한다.",
        "각자 테스트 레벨에서 먼저 작업하고, 통합 시점에 메인 레벨로 옮긴다.",
        "블루프린트 이름과 폴더 구조를 통일한다.",
        "공통 부모 클래스나 공통 인터페이스는 이진헌이 관리한다.",
        "코어 로직, 인터랙션, 연출, 콘텐츠 폴더를 분리한다.",
    ])

    doc.add_heading("6. 주차별 진행", level=1)
    add_matrix(
        doc,
        ["시기", "목표", "작업"],
        [
            ["Week 1", "회색박스 코어와 대표 분위기", "코어 루프, ArtToneTest, InteractableBase, 에셋 리스트"],
            ["Week 2", "사물형 UI와 기본 연출 연결", "AI/스테이지, 룰렛/생명 카드, LLM 테스트, 핵심 소품 배치"],
            ["Week 3", "기술 쇼케이스 성공", "DB, 멀티 테스트맵, Stage별 분위기, 스토리 소품"],
            ["Week 4", "통합 빌드", "버그 수정, 플레이테스트, 연출 타이밍 조정"],
            ["마무리", "평가 시연 안정화", "시연 순서 고정, 영상/스크린샷 확보, 실패 대비 플랜"],
        ],
        [1200, 2500, 5660],
    )

    doc.add_heading("7. 최종 체크리스트", level=1)
    add_bullets(doc, [
        "한 명이 하루 이상 막히면 바로 공유한다.",
        "복잡한 기능은 테스트 레벨에서 먼저 성공시킨다.",
        "코어 로직과 연출은 이벤트로만 연결한다.",
        "매주 최소 한 번은 develop에서 전체가 플레이 가능한 빌드를 만든다.",
        "평가 전에는 새 기능 추가를 멈추고 이미 되는 기능을 안정화한다.",
        "시연에서는 기능을 왜 넣었는지 설명한다: 사물형 UI는 몰입감, LLM은 살아있는 보스, DB는 기록 저장, 멀티는 네트워크 동기화.",
    ])
    return save_doc(doc, "ShowDown_Collaboration_Common.docx")


def main():
    paths = [
        build_im(),
        build_lee(),
        build_park(),
        build_kim(),
        build_collab(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()

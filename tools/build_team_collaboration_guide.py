from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ShowDown_Team_Collaboration_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    for line in lines:
        pp = cell.add_paragraph(style=None)
        pp.add_run(line)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], "E8EEF5")
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.add_run(value)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Malgun Gothic"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.font.bold = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup_styles(doc)

    doc.add_paragraph("ShowDown", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.add_run("4인 초보 팀 협업 가이드: 역할별 코드 경계, 주의점, 실전 팁").bold = True
    doc.add_paragraph("기준 역할: 임현진(연출 메인), 이진헌(기술 통합/연출 보조/멀티/LLM/DB/Git), 박형빈(코어 게임 루프), 김윤아(에셋/맵/스토리 소품)")
    doc.add_paragraph("목표: 단순한 카드 룰을 강한 사물형 UI와 룰렛 연출, 기술 쇼케이스 기능으로 평가자에게 선명하게 보여주는 것.")

    add_callout(doc, "핵심 결론", [
        "코어 로직, 사물형 UI, 연출, 에셋/맵은 서로 직접 얽히지 말고 이벤트로 연결한다.",
        "박형빈은 게임 규칙을 맞게 돌리고, 임현진은 그 결과가 강하게 보이게 만들고, 이진헌은 둘을 안정적으로 연결하고, 김윤아는 공간과 소품의 설득력을 만든다.",
        "모든 기능은 먼저 작게 성공시킨 뒤 본 게임에 붙인다. 특히 멀티, LLM, DB는 평가용 쇼케이스 범위로 시작한다.",
    ])

    doc.add_heading("1. 전체 협업 원칙", level=1)
    add_bullets(doc, [
        "같은 파일을 동시에 만지지 않는다. Unreal의 .umap, .uasset 충돌은 초보 팀에게 거의 수동 복구가 어렵다.",
        "코어 게임 로직은 UI와 연출을 몰라야 한다. 코어는 결과 이벤트만 발생시키고, 연출은 그 이벤트를 받아 반응한다.",
        "처음부터 예쁜 최종물을 만들지 않는다. 회색박스 코어, 테스트맵 연출, 테스트 API, 테스트 DB처럼 각 기능을 분리해서 먼저 성공시킨다.",
        "기능 구현 담당자는 '돌아간다'를 책임지고, 연출 담당자는 '플레이어가 이해하고 느낀다'를 책임진다.",
        "완성된 것만 합친다. 각자 브랜치에서 깨진 상태로 오래 두지 말고, 작게 커밋하고 자주 공유한다.",
    ])

    doc.add_heading("2. 추천 코드 구조", level=1)
    doc.add_paragraph("코드는 크게 네 층으로 나눈다. 초보 팀일수록 이 경계가 중요하다.")
    add_matrix(
        doc,
        ["층", "책임", "예시"],
        [
            ["Core", "게임 규칙과 상태만 처리", "덱, 손패, 베팅, 승패, 룰렛 확률, 목숨, 스테이지"],
            ["Interaction", "플레이어가 사물을 클릭하거나 조작하는 입력 처리", "테이블 버튼, 총알 선택 장치, 카드 선택 오브젝트"],
            ["Presentation", "카메라, 사운드, VFX, UI 피드백 재생", "룰렛 시퀀스, 생명 카드 소실, 화면 노이즈, 조명 깜빡임"],
            ["Content", "맵, 소품, 텍스트, 에셋 배치", "방 드레싱, CRT, 벽 낙서, 보스 마스크, 스토리 단서"],
        ],
        [1500, 2500, 5360],
    )

    doc.add_heading("3. 이벤트 기반 연결 규칙", level=1)
    doc.add_paragraph("아래 이벤트 이름을 팀 공용 약속으로 두면 협업이 쉬워진다. 실제 C++ Delegate, Blueprint Event Dispatcher, Interface 중 어떤 방식이든 좋지만 이름과 데이터는 통일한다.")
    add_matrix(
        doc,
        ["이벤트", "누가 발생", "누가 사용"],
        [
            ["OnPhaseChanged(NewPhase)", "코어", "연출, 사물 UI, 모니터 UI"],
            ["OnCardSelected(Player, CardId)", "코어", "카드 오브젝트, 카드 부착 연출"],
            ["OnBetChanged(Player, BulletCount)", "코어", "총알 오브젝트, 베팅 장치, CRT"],
            ["OnActionResolved(ActionType)", "코어", "Call/Raise/Fold 버튼 피드백"],
            ["OnCardsRevealed(PlayerCard, BossCard)", "코어", "카드 공개 카메라, 사운드"],
            ["OnRoundResolved(Result)", "코어", "승패/동점 연출, 룰렛 시작 준비"],
            ["OnRouletteStarted(Target, BulletCount)", "코어", "리볼버 장전 시퀀스"],
            ["OnRouletteResult(Target, bHit)", "코어", "명중/불발 연출"],
            ["OnLifeChanged(Target, Life)", "코어", "생명 카드 소실/갱신"],
            ["OnStageChanged(Stage)", "코어", "방 분위기, 보스 변질, Post Process"],
        ],
        [2450, 1650, 5260],
    )

    doc.add_paragraph("금지 규칙: 코어 함수 안에서 카메라 흔들림, 사운드 재생, 총알 애니메이션을 직접 실행하지 않는다. 반대로 연출 블루프린트 안에서 승패 계산, 카드 확률, 룰렛 명중 판정을 하지 않는다.")

    doc.add_heading("4. 임현진: 연출 메인 / 사물형 UI / VFX", level=1)
    doc.add_paragraph("임현진은 게임의 첫인상과 기억에 남는 순간을 책임진다. 코어 로직을 많이 건드리기보다, 게임 규칙이 만든 사건을 플레이어가 강하게 느끼도록 만든다.")

    doc.add_heading("4.1 먼저 만들 것", level=2)
    add_numbered(doc, [
        "LV_ArtToneTest: 어두운 방, 테이블, 보스 실루엣, 카드, 총, 총알, 생명 카드를 배치한 대표 장면.",
        "LV_EffectTest: 버튼으로 카드 공개, 총알 장전, 방아쇠, 명중/불발, 생명 카드 소실을 각각 테스트하는 맵.",
        "BP_PresentationDirector: 코어 이벤트를 받아 어떤 연출을 재생할지 관리하는 중앙 연출 관리자.",
        "BP_RouletteSequence: BulletCount와 Hit/Miss를 입력받아 장전, 약실 회전, 방아쇠, 결과 연출을 재생.",
        "BP_LifeCardDisplay: 목숨 수에 따라 테이블 위 생명 카드가 사라지거나 오염되는 연출.",
        "PP_AnalogHorror_Base: 기본 포스트프로세스 프리셋. 낮은 채도, 높은 대비, vignette, film grain, 약한 chromatic aberration.",
    ])

    doc.add_heading("4.2 코드/블루프린트 작성 방식", level=2)
    add_bullets(doc, [
        "연출 블루프린트는 입력값을 받는 함수 형태로 만든다. 예: PlayRouletteSequence(BulletCount, bHit), PlayLifeLost(Target, NewLife), PlayRevealCards(PlayerCard, BossCard).",
        "테스트 버튼으로도 실행되고, 나중에 코어 이벤트로도 실행되게 만든다. 이러면 코어 루프를 기다리지 않고 작업할 수 있다.",
        "연출 중에는 입력을 잠그는 상태가 필요하다. BP_PresentationDirector에 bIsPlayingPresentation 같은 변수를 두고, 끝나면 입력을 다시 연다.",
        "카메라, 사운드, VFX 타이밍은 한 블루프린트에 흩뿌리지 말고 Sequence 단위로 묶는다.",
        "값을 하드코딩하지 말고 DataAsset 또는 변수로 뺀다. 예: Stage별 vignette 강도, flicker 빈도, 노이즈 강도.",
    ])

    doc.add_heading("4.3 주의점", level=2)
    add_bullets(doc, [
        "모든 UI를 사물형으로 만들면 멋있지만, 시간이 터진다. 우선 플레이 중 핵심 UI만 사물화한다: 체력, 베팅, Call/Raise/Fold, 룰렛 위험도.",
        "글씨를 완전히 없애려고 하다가 플레이어가 룰을 이해 못 하면 실패다. 필요하면 CRT 안의 짧은 글자, 카드 라벨, 버튼 각인처럼 게임 안 사물에 최소 텍스트를 넣는다.",
        "연출이 길면 반복 플레이가 피곤하다. 첫 등장만 길고, 반복 시 짧아지는 옵션을 고려한다.",
        "어둡게 만드는 것과 안 보이게 만드는 것은 다르다. 카드, 총알, 리볼버, 생명 카드는 항상 읽혀야 한다.",
        "Post Process를 과하게 걸면 UI와 카드 숫자가 안 보인다. Stage 3 전용 강한 효과와 기본 플레이 효과를 분리한다.",
    ])

    doc.add_heading("4.4 평가 어필 포인트", level=2)
    add_bullets(doc, [
        "2D UI를 최소화하고, 게임 안 사물로 상태와 선택을 표현했다.",
        "룰렛 장면의 카메라, 사운드, 조명, VFX를 한 시퀀스로 설계했다.",
        "단순한 카드 게임 루프를 사물형 UI와 연출로 긴장감 있게 확장했다.",
    ])

    doc.add_heading("5. 이진헌: 기술 통합 / 연출 보조 / 멀티 / LLM / DB / Git", level=1)
    doc.add_paragraph("이진헌은 팀의 기술 연결점이다. 직접 모든 것을 완성형으로 크게 만들기보다, 각 기술을 작고 확실하게 성공시킨 뒤 본 게임에 붙이는 것이 핵심이다.")

    doc.add_heading("5.1 먼저 만들 것", level=2)
    add_numbered(doc, [
        "Git 브랜치 구조와 폴더 소유권 정리. develop 브랜치에 안정적으로 통합한다.",
        "BP_InteractableBase: 모든 클릭 가능한 사물의 공통 부모. Hover, Unhover, Interact, Enable/Disable을 통일.",
        "BP_DiegeticButton: 테이블 버튼, 스위치, 레버 등 사물형 버튼의 공통 구현.",
        "BP_BulletSelector: 1~6발 선택 장치. 선택하면 Core의 RaiseTo(BulletCount)를 호출.",
        "LLM_TestLevel: 본 게임과 분리해서 API 응답 받기만 먼저 성공.",
        "DB_TestWidget 또는 TestLevel: 기억 조각, 승패 기록, 스킨 해금 중 하나를 저장/불러오기.",
        "Multiplayer_TestMap: 두 클라이언트 접속, 간단한 상태 동기화, 버튼 입력 복제부터 검증.",
    ])

    doc.add_heading("5.2 코드 작성 방식", level=2)
    add_bullets(doc, [
        "Interactable은 코어 규칙을 직접 계산하지 않는다. 클릭되면 GameManager의 공개 함수만 호출한다.",
        "서버/호스트 권위가 필요한 멀티 기능은 처음부터 '누가 최종 결정을 내리는가'를 정한다. 카드 지급, 승패, 룰렛 결과는 클라이언트 마음대로 계산하면 안 된다.",
        "LLM은 게임 판단을 맡기지 않는다. 보스의 실제 행동은 코드가 결정하고, LLM은 대사만 생성한다.",
        "LLM 호출은 비동기로 처리한다. 응답이 늦으면 대사 풀에서 즉시 대체한다.",
        "DB는 처음부터 복잡한 계정 시스템을 만들지 않는다. 평가용이면 플레이 기록, 재화, 스킨 해금, 간단 랭킹 중 하나를 확실히 저장한다.",
        "Git 병합 전에는 반드시 에디터에서 열리는지, 컴파일 되는지, 주요 테스트맵이 깨지지 않는지 확인한다.",
    ])

    doc.add_heading("5.3 주의점", level=2)
    add_bullets(doc, [
        "Git 담당자가 모든 충돌을 혼자 해결하려고 하면 위험하다. 충돌이 난 파일의 담당자를 불러서 같이 본다.",
        "멀티, LLM, DB를 한 번에 본 게임에 붙이면 어디서 깨졌는지 모른다. 반드시 테스트 레벨에서 따로 성공시킨다.",
        "API Key를 GitHub에 올리면 안 된다. Config 파일이나 환경변수, 로컬 파일로 분리하고 .gitignore에 넣는다.",
        "멀티는 범위를 줄인다. 목표는 사설 방 1:1에서 핵심 상태가 동기화되는 것. 랭크, 시즌, 복잡한 매칭은 후순위다.",
        "Unreal Asset 충돌은 텍스트 코드 충돌보다 훨씬 아프다. 통합 담당자는 공용 맵과 공용 블루프린트 수정 시간을 관리한다.",
    ])

    doc.add_heading("5.4 평가 어필 포인트", level=2)
    add_bullets(doc, [
        "사물형 UI 인터랙션 시스템을 공통화했다.",
        "LLM은 게임을 멈추지 않도록 비동기와 폴백 구조로 붙였다.",
        "DB 저장/불러오기 또는 랭킹을 통해 외부 데이터 연동을 보여줬다.",
        "멀티는 1:1 상태 동기화와 서버 권위 구조를 보여줬다.",
    ])

    doc.add_heading("6. 박형빈: 코어 게임 루프 / 룰 시스템 / 기본 보스 AI", level=1)
    doc.add_paragraph("박형빈은 게임이 규칙대로 돌아가게 만드는 역할이다. 화면이 예쁘지 않아도 상관없지만, 한 판이 끝까지 안정적으로 돌아가야 한다.")

    doc.add_heading("6.1 먼저 만들 것", level=2)
    add_numbered(doc, [
        "BP_GameManager 또는 C++ GameManager: 현재 Phase, 턴, 목숨, 스테이지 상태를 관리.",
        "CardSystem: 1~7 각 2장, 총 14장 덱 생성, 각자 5장 지급, 4장 미사용 처리.",
        "BettingSystem: Check, RaiseTo, Call, Fold 처리. 한 판 최대 6발 제한.",
        "RoundResolver: 카드 공개, 높은 숫자 승리, 동점 양쪽 룰렛 처리.",
        "RouletteSystem: BulletCount / 6 확률로 Hit/Miss 결정. 6발은 100% 명중.",
        "LifeSystem: 목숨 감소, 사망, 스테이지 클리어, 게임 오버 처리.",
        "BossAI_Basic: Stage별로 보수적/공격적 베팅 성향을 바꾸는 간단 AI.",
    ])

    doc.add_heading("6.2 공개 함수 예시", level=2)
    add_bullets(doc, [
        "SelectCard(PlayerId, CardId): 선택 가능한 카드인지 검사하고 상대 머리 위 카드로 등록.",
        "RaiseTo(PlayerId, BulletCount): 현재 베팅보다 큰지, 1~6 범위인지 검사.",
        "Call(PlayerId): 상대 최종 베팅 수와 맞춘다.",
        "Fold(PlayerId): 그 판 패배 처리. 7 폴드 예외면 6발 강제.",
        "RevealCards(): 양쪽 카드 공개 후 승패를 판정.",
        "ResolveRoulette(TargetPlayer, BulletCount): 명중 여부를 계산하고 이벤트를 발생.",
        "StartNextRound(): 카드 폐기, 다음 카드 선택 단계로 전환.",
    ])

    doc.add_heading("6.3 주의점", level=2)
    add_bullets(doc, [
        "처음부터 사물형 UI에 맞춰 만들지 말고, 임시 2D 버튼이나 키보드 입력으로도 코어 루프가 돌아가게 만든다.",
        "코어 로직 안에서 연출을 직접 부르지 않는다. 상태가 바뀌면 이벤트만 발생시킨다.",
        "카드 상태를 명확히 나눈다: 손패, 머리 위 카드, 공개된 카드, 버려진 카드, 미사용 카드.",
        "한 함수가 너무 많은 일을 하지 않게 한다. 예: Fold는 폴드 처리만 하고, 룰렛 연출은 이벤트 이후 Presentation 쪽에서 처리.",
        "디버그 로그를 많이 남긴다. 초보 팀에서는 현재 Phase, 현재 베팅, 카드 값, 룰렛 결과 로그가 버그 찾기에 큰 도움이 된다.",
        "AI는 처음엔 단순하게 만든다. 완벽한 심리전 AI보다 Stage별 성향이 보이는 AI가 먼저다.",
    ])

    doc.add_heading("6.4 완료 기준", level=2)
    add_bullets(doc, [
        "아트 없이도 5판 이상 연속 진행 가능.",
        "카드 선택, 베팅, 공개, 룰렛, 목숨 감소, 다음 판이 끊기지 않음.",
        "7 폴드 예외, 동점 양쪽 룰렛, 6발 확정 명중이 테스트됨.",
        "연출 담당자가 받을 수 있는 이벤트가 모두 발생함.",
    ])

    doc.add_heading("7. 김윤아: 에셋 / 맵 드레싱 / 스토리 소품", level=1)
    doc.add_paragraph("김윤아는 코딩 부담을 줄이는 대신, 게임 공간이 설득력 있게 보이도록 만드는 역할이다. 단순히 에셋을 찾는 것에서 끝나지 않고, 실제 방의 분위기와 소품 배치까지 책임지는 것이 좋다.")

    doc.add_heading("7.1 먼저 만들 것", level=2)
    add_numbered(doc, [
        "에셋 리스트 스프레드시트: 이름, 용도, 출처, 라이선스, 다운로드 링크, 사용 위치.",
        "방 콘셉트 보드: Buckshot Roulette 계열의 좁은 방, 낡은 테이블, 아날로그 호러, CRT, 후드 보스 레퍼런스.",
        "LV_Environment_Blockout: 테이블, 보스 자리, 플레이어 자리, 조명, 카메라 위치가 잡힌 맵 초안.",
        "핵심 소품 세트: 카드, 리볼버, 총알/코인, 생명 카드, 테이블, 조명, CRT.",
        "스토리 단서 소품: 벽 긁힌 글자, 타다 만 카드, 메모, 이전 희생자의 흔적.",
        "Stage별 변화안: Stage 1은 정돈, Stage 2는 오염, Stage 3은 붕괴와 왜곡.",
    ])

    doc.add_heading("7.2 작업 방식", level=2)
    add_bullets(doc, [
        "에셋은 무료라고 다 써도 되는 것이 아니다. 상업/비상업 가능 여부, 출처 표기 필요 여부를 기록한다.",
        "폴더를 정리한다. 예: Content/ShowDown/Environment, Props, Characters, Materials, UIObjects, VFX.",
        "소품은 게임 정보와 연결되는 것을 우선한다. 예쁜 배경 소품보다 카드, 총, 총알, 생명 카드, 버튼 장치가 먼저다.",
        "맵 드레싱은 플레이 가독성을 해치면 안 된다. 카드와 총알을 보는 데 방해되는 소품은 줄인다.",
        "보스 마스크는 실제 유명인 얼굴을 그대로 쓰지 말고, '유명인풍 부틀렉 마스크'처럼 변형된 오리지널 방향으로 잡는다.",
    ])

    doc.add_heading("7.3 주의점", level=2)
    add_bullets(doc, [
        "메인 맵 파일은 한 번에 한 명만 수정한다. 맵 배치 작업 시간대를 정하고, 작업 전후에 팀에게 알린다.",
        "너무 많은 고해상도 에셋을 넣으면 프로젝트가 무거워진다. 평가용 핵심 화면에 보이는 것부터 넣는다.",
        "에셋 이름을 대충 두면 나중에 찾기 어렵다. SM_Table_Rusty, M_Card_Burnt, BP_CRT_Monitor처럼 역할이 보이게 짓는다.",
        "텍스처/머티리얼 색감은 임현진의 Post Process와 함께 확인한다. 에디터 뷰에서 예뻐도 실제 플레이 카메라에서 안 보이면 수정한다.",
    ])

    doc.add_heading("8. Git 협업 규칙", level=1)
    add_matrix(
        doc,
        ["브랜치", "담당", "용도"],
        [
            ["main", "이진헌", "시연 가능한 안정 버전만 유지"],
            ["develop", "이진헌", "통합 브랜치. 각 기능이 여기로 모임"],
            ["feature/core-loop", "박형빈", "싱글 코어 루프와 룰 시스템"],
            ["feature/presentation", "임현진", "연출, VFX, 사물형 UI 콘셉트"],
            ["feature/tech-integration", "이진헌", "인터랙션, LLM, DB, 멀티, 통합"],
            ["feature/environment", "김윤아", "맵, 에셋, 소품 배치"],
        ],
        [2300, 1600, 5460],
    )
    add_bullets(doc, [
        "커밋은 작게 한다. '룰렛 다 함'보다 'RouletteSystem 명중 판정 추가'처럼 추적 가능하게 쓴다.",
        "커밋 전에 에디터가 열리는지 확인한다.",
        "공용 파일을 수정해야 하면 먼저 단체 채팅에 말한다.",
        "충돌이 나면 담당자 없이 억지로 해결하지 않는다.",
        "Binaries, Intermediate, Saved, DerivedDataCache는 Git에 올리지 않는다.",
        "API Key, DB 비밀번호, 개인 설정 파일은 절대 커밋하지 않는다.",
    ])

    doc.add_heading("9. 주차별 추천 진행", level=1)
    add_matrix(
        doc,
        ["시기", "목표", "각자 할 일"],
        [
            ["Week 1", "한 판이 도는 회색박스와 대표 분위기 확보", "박형빈: 코어 루프 / 임현진: ArtToneTest / 이진헌: Git+InteractableBase / 김윤아: 에셋 리스트+맵 블록아웃"],
            ["Week 2", "사물형 UI와 기본 연출 연결", "박형빈: AI+스테이지 / 임현진: 룰렛+생명 카드 / 이진헌: UI 연결+LLM 테스트 / 김윤아: 핵심 소품 배치"],
            ["Week 3", "기술 쇼케이스 분리 성공", "이진헌: DB, 멀티 테스트맵 / 임현진: Stage별 분위기 / 박형빈: 밸런싱 이벤트 / 김윤아: 스토리 소품"],
            ["Week 4", "통합 빌드", "전원: 버그 수정, 플레이테스트, 연출 타이밍 조정"],
            ["마무리", "평가 시연용 안정화", "시연 순서 고정, 영상/스크린샷 확보, 실패 대비 플랜 준비"],
        ],
        [1200, 2500, 5660],
    )

    doc.add_heading("10. 초보 팀을 위한 최종 체크리스트", level=1)
    add_bullets(doc, [
        "한 명이 하루 이상 막히면 바로 공유한다. 초보 팀에서 침묵은 가장 큰 리스크다.",
        "복잡한 기능은 테스트 레벨에서 먼저 성공시킨다.",
        "사물형 UI는 핵심 플레이에 필요한 것부터 만든다.",
        "코어 로직과 연출은 이벤트로만 연결한다.",
        "매주 최소 한 번은 develop에서 전체가 플레이 가능한 빌드를 만든다.",
        "평가 전에는 새 기능 추가를 멈추고, 이미 되는 기능의 안정성과 시연성을 올린다.",
        "시연에서는 '이 기능을 왜 넣었는지'까지 말한다: 사물형 UI는 몰입감, LLM은 살아있는 보스, DB는 기록 저장, 멀티는 네트워크 동기화.",
    ])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("ShowDown Team Collaboration Guide")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
